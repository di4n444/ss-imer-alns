"""Native Word equations (OMML) for the thesis.

The faculty template requires mathematical expressions as real Word equation objects with
each symbol formatted properly, not as pictures or italicised plain text. Word's own
equations are OMML - Office Math Markup Language - so this module builds OMML elements
directly and python-docx inserts them into the document.

Two things follow from OMML being real math markup rather than styled text:

  * A letter inside `m:r` is rendered by Word in math italic automatically, which is the
    correct convention for a variable. Anything that must stay upright - a function name
    like exp, an operator name like arg min, a word used as a label - carries `m:nor`
    ("normal text") and is written with `up()`.
  * Structure is structural: a fraction is `m:f` with numerator and denominator, a
    subscript is `m:sSub`, a sum is `m:nary` with its own limits. Word can therefore
    reflow, edit and re-render them, and they scale with the surrounding text.

Numbering follows the template, which places each numbered expression in a two-cell
borderless table: the equation centred in the first cell, the number right-aligned in the
second. Reference them in the text as "prema izrazu (1)".
"""

from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Matches the run properties Word itself writes inside an equation, and the template's
# own: Cambria Math at 12 pt (w:sz is in half-points).
_MATH_FONT = "Cambria Math"
_HALF_POINTS = "24"


def _el(tag: str, *children, **attrs):
    node = OxmlElement(tag)
    for key, value in attrs.items():
        node.set(qn(key.replace("_", ":")), value)
    for child in children:
        if child is not None:
            node.append(child)
    return node


def _run_props(upright: bool = False):
    fonts = _el("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        fonts.set(qn(attr), _MATH_FONT)
    props = _el("w:rPr", fonts, _el("w:sz", w_val=_HALF_POINTS),
                _el("w:szCs", w_val=_HALF_POINTS))
    if upright:
        props.append(_el("w:i", w_val="0"))
        props.append(_el("w:iCs", w_val="0"))
    return props


def _ctrl():
    """Control properties - Word attaches these to every structured element."""
    return _el("m:ctrlPr", _run_props())


def _text(text: str):
    node = OxmlElement("m:t")
    node.set(qn("xml:space"), "preserve")
    node.text = text
    return node


def _mk_run(text: str, upright: bool):
    """Math style is set explicitly rather than left to Word's default.

    `m:sty` is the reliable control: "p" is plain upright, "i" is math italic. `m:nor`
    alone was not enough - function names still rendered slanted - so the character-level
    italic flag is switched off alongside it for upright runs."""
    if upright:
        props = _el("m:rPr", _el("m:sty", m_val="p"), _el("m:nor"))
        return _el("m:r", props, _run_props(upright=True), _text(text))
    return _el("m:r", _el("m:rPr", _el("m:sty", m_val="i")), _run_props(), _text(text))


def v(text: str):
    """Variable or symbol: rendered in math italic by Word."""
    return _mk_run(text, upright=False)


def up(text: str):
    """Upright text: function names (exp, ln), operator names (arg min), and any word
    used as a label. Keeping these upright is the difference between a formatted formula
    and italicised prose."""
    return _mk_run(text, upright=True)


def row(*items):
    """Several elements in sequence, e.g. v('a'), up(' + '), v('b')."""
    return list(items)


def _wrap(items):
    """Accept an element, or any nesting of lists of them, and return one flat list.

    Nesting arises naturally because helpers like `func` return several elements, and
    those get passed straight into another helper's argument slot."""
    if items is None:
        return []
    if not isinstance(items, (list, tuple)):
        return [items]
    flat = []
    for item in items:
        flat.extend(_wrap(item))
    return flat


def sub(base, subscript):
    """Subscript, e.g. p_uv."""
    return _el("m:sSub",
               _el("m:sSubPr", _ctrl()),
               _el("m:e", *_wrap(base)),
               _el("m:sub", *_wrap(subscript)))


def sup(base, superscript):
    """Superscript, e.g. 2^|E|."""
    return _el("m:sSup",
               _el("m:sSupPr", _ctrl()),
               _el("m:e", *_wrap(base)),
               _el("m:sup", *_wrap(superscript)))


def frac(numerator, denominator):
    """Built-up fraction."""
    return _el("m:f",
               _el("m:fPr", _ctrl()),
               _el("m:num", *_wrap(numerator)),
               _el("m:den", *_wrap(denominator)))


def delim(*items, left="(", right=")"):
    """Delimiters that grow with their content - parentheses, brackets, braces, and
    with left='|' right='|' the absolute value or cardinality bars."""
    props = _el("m:dPr", _el("m:begChr", m_val=left), _el("m:endChr", m_val=right), _ctrl())
    return _el("m:d", props, _el("m:e", *_wrap(list(items))))


def nary(operator, lower=None, upper=None, body=None, hide_upper=False):
    """n-ary operator with limits: sums, products, unions."""
    props = _el(
        "m:naryPr",
        _el("m:chr", m_val=operator),
        _el("m:limLoc", m_val="undOvr"),
        _el("m:subHide", m_val="1") if lower is None else None,
        _el("m:supHide", m_val="1") if (upper is None or hide_upper) else None,
        _ctrl(),
    )
    return _el("m:nary", props,
               _el("m:sub", *_wrap(lower)),
               _el("m:sup", *_wrap(upper)),
               _el("m:e", *_wrap(body)))


def limlow(base, limit):
    """A limit set underneath an operator, which is where the constraint of an arg min
    belongs in a displayed expression. Written as a subscript it would sit to the right
    and read as an index instead."""
    return _el("m:limLow",
               _el("m:limLowPr", _ctrl()),
               _el("m:e", *_wrap(base)),
               _el("m:lim", *_wrap(limit)))


def acc(base, char="̂"):
    """An accent over a symbol - by default the hat that marks an estimator, so that
    sigma-hat is one properly built object rather than a letter followed by a combining
    character that may or may not compose in the reader's font."""
    return _el("m:acc",
               _el("m:accPr", _el("m:chr", m_val=char), _ctrl()),
               _el("m:e", *_wrap(base)))


def func(name, argument):
    """A named function applied to an argument, e.g. exp(x): the name upright, the
    argument in its own delimiters."""
    return [up(name), delim(argument)]


def omath(*items):
    """Assemble elements into one inline equation object."""
    parts = []
    for item in items:
        parts.extend(_wrap(item))
    return _el("m:oMath", *parts)


# -- placement ---------------------------------------------------------------

def _borderless(table):
    """The template's numbered-equation table shows no rules; python-docx's default
    table style would."""
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        borders.append(_el(f"w:{edge}", w_val="none", w_sz="0", w_space="0"))
    table._tbl.tblPr.append(borders)


def add_equation(document, *items, number=None, width_cm=15.0):
    """Insert a display equation, numbered as the template requires: the expression
    centred, the number right-aligned against the margin."""
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm

    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _borderless(table)

    body, label = table.rows[0].cells
    body.width, label.width = Cm(width_cm * 0.88), Cm(width_cm * 0.12)

    paragraph = body.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph._p.append(omath(*items))

    label_paragraph = label.paragraphs[0]
    label_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if number is not None:
        label_paragraph.add_run(f"({number})")
    return table


def add_inline(paragraph, *items):
    """An equation inside a sentence, for symbols mentioned in running text."""
    paragraph._p.append(omath(*items))


class Italic(str):
    """Plain italic text: a foreign term, a problem name, a book title.

    These are words, not mathematics, so they must not become equation objects - an
    English term set as an OMML run is typeset in a maths font and cannot be spell-checked
    or edited as text."""


def i(text: str) -> "Italic":
    return Italic(text)


def para(container, *parts, style=None):
    """A paragraph mixing prose and mathematics.

    Plain strings become ordinary runs; anything else is inserted as an inline equation
    object. Every symbol named in the running text therefore carries the same formatting
    as in the displayed expressions, which is what the template asks for - writing a
    variable as an ordinary italic letter in prose and as an equation object in a
    displayed formula would set the same symbol in two different typefaces.

        para(doc, "Neka je ", v("G"), " = ", delim(v("V"), up(", "), v("E")), " graf.")
    """
    paragraph = container.add_paragraph(style=style)
    for part in parts:
        if isinstance(part, Italic):
            paragraph.add_run(str(part)).italic = True
        elif isinstance(part, str):
            paragraph.add_run(part)
        else:
            add_inline(paragraph, part)
    return paragraph
