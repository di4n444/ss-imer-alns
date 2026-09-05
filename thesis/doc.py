"""Document assembly against the faculty template's rules.

The template fixes the things this module enforces so no chapter has to remember them:

  * body text is Times New Roman 12 pt with 1.5 line spacing, on A4;
  * mathematical expressions are numbered "(n)" against the right margin and referenced
    in the text as "prema izrazu (n)" - numbering is continuous through the document;
  * figures are numbered <poglavlje>.<redni broj u poglavlju>, captioned *below* as
    "Sl. 2.1 ..." and referenced in the text before they appear;
  * tables carry the same numbering but are captioned *above*, as "Tablica 2.1 ...";
  * code listings use the "Kôd" style and are captioned below as "Kôd 5.1 - ...";
  * lists use the template's own "bullet1" style rather than a hand-written marker;
  * headings use the template's own Heading 1/2/3 styles, so Word's table of contents and
    the numbering scheme keep working, and go no deeper than level 3, which the template
    explicitly rules out.

A `Thesis` instance carries the counters, so a chapter writes `t.eq(...)` or
`t.figure(...)` without tracking numbers by hand and without them going stale when a
section moves.
"""

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt

import bibliography as bib
import omml as M


class Thesis:
    """Writes chapters into an existing document, immediately before `marker_heading`.

    python-docx only appends to the end of the body, so every block is created and then
    relocated in front of the marker. Doing that per block, rather than moving a batch
    afterwards, means the document is in a valid order at every point and content can
    never be left stranded after the bibliography."""

    def __init__(self, document, marker_heading: str = "Literatura", numbering=None,
                 sections=None):
        self.d = document
        self.numbering = numbering or {}
        self.chapter = 0
        self._equation = 0
        self._figures = {}
        self._tables = {}
        self._listings = {}
        self._level2 = 0
        self._level3 = 0
        self.refs = {}          # label -> number, so cross-references cannot drift
        # Section numbers resolved on a previous pass. A chapter may refer forward to a
        # section that has not been written yet (chapter 3 points at 5.2.4), so the number
        # cannot be known while the first pass is running - see `sec`.
        self.sections = {}
        self._known_sections = sections or {}
        self.marker = None
        for p in document.paragraphs:
            if p.style.name == "Heading 1" and p.text.strip() == marker_heading:
                self.marker = p._p
                break
        if self.marker is None:
            raise ValueError(f"heading '{marker_heading}' not found")

    def _place(self, block):
        """Move a freshly appended block into position before the marker."""
        element = block._p if hasattr(block, "_p") else block._tbl
        self.marker.addprevious(element)
        return block

    # -- structure --------------------------------------------------------

    def _heading(self, text: str, level: int, numbered: bool):
        """Chapter numbering in this template is direct paragraph formatting, not part of
        the Heading style, so a generated heading only joins the numbering sequence if it
        carries the same `numPr` as the template's own headings. Uvod, Zaključak and
        Literatura deliberately go without."""
        paragraph = self.d.add_paragraph(text, style=f"Heading {level}")
        if not numbered:
            _suppress_numbering(paragraph)
        return self._place(paragraph)

    def h1(self, text: str, numbered: bool = True, label: str = None):
        if numbered:
            self.chapter += 1
            self._figures[self.chapter] = 0
            self._tables[self.chapter] = 0
            self._listings[self.chapter] = 0
            self._level2 = self._level3 = 0
            self._record(label, str(self.chapter))
        return self._heading(text, 1, numbered)

    def h2(self, text: str, label: str = None):
        self._level2 += 1
        self._level3 = 0
        self._record(label, f"{self.chapter}.{self._level2}")
        return self._heading(text, 2, True)

    def h3(self, text: str, label: str = None):
        self._level3 += 1
        self._record(label, f"{self.chapter}.{self._level2}.{self._level3}")
        return self._heading(text, 3, True)

    def _record(self, label: str, number: str):
        if label:
            self.sections[label] = number

    def sec(self, label: str) -> str:
        """The number of a labelled section, e.g. "5.2.4".

        Word numbers the headings itself, so the number exists only as a position in the
        document and cannot be read back from the text. It is therefore collected on a
        first pass and substituted on a second — which is what makes a forward reference
        possible at all, and what stops every reference in the document from silently
        pointing one section off when a section is inserted."""
        try:
            return self._known_sections[label]
        except KeyError:
            # First pass: the number is not known yet. Returning a marker rather than
            # raising lets the pass finish and collect every label; build_thesis refuses
            # to save a document that still contains one.
            return "?.?"

    # -- body -------------------------------------------------------------

    def p(self, *parts):
        """A paragraph of prose, optionally with inline mathematics. Strings are text -
        with "{key}" citations resolved against the bibliography - and anything else is
        an equation object."""
        return self._place(M.para(self.d, *self._cited(parts)))

    @staticmethod
    def _cited(parts):
        """Resolve citations inside text, preserving whether a part is italic."""
        out = []
        for part in parts:
            if isinstance(part, M.Italic):
                out.append(M.Italic(bib.expand(str(part))))
            elif isinstance(part, str):
                out.append(bib.expand(part))
            else:
                out.append(part)
        return out

    def bullets(self, items, numbered: bool = False):
        """A list.

        Unnumbered lists use the template's own "bullet1" style, which supplies the
        bullet glyph itself - so no marker is written into the text. Numbered lists keep a
        hand-written "1." marker: the template's "bullet1brojevi" numbers items as "[1]",
        which reads as a citation rather than as a step, and these lists enumerate the
        rules of a model."""
        names = self._style_names()
        if numbered:
            style = "Normal Indent" if "Normal Indent" in names else None
        else:
            style = "bullet1" if "bullet1" in names else None

        for n, item in enumerate(items, start=1):
            parts = item if isinstance(item, (list, tuple)) else [item]
            prefix = [f"{n}. "] if numbered else []
            self._place(M.para(self.d, *prefix, *self._cited(parts), style=style))

    def _style_names(self):
        return {s.name for s in self.d.styles}

    def code(self, lines, caption: str, label: str = None):
        """A code or pseudocode listing in the template's "Kôd" style, numbered like a
        figure and captioned below it as "Kôd <poglavlje>.<n> - ...".

        Line breaks are soft (`add_break`) rather than one paragraph per line, so the
        whole listing stays a single block that Word will not split mid-listing or
        re-space as if each line were a paragraph."""
        number = self._next_listing_number(label)
        style = "Kôd" if "Kôd" in self._style_names() else None

        block = self._place(self.d.add_paragraph(style=style))
        run = block.add_run()
        for i, line in enumerate(lines):
            if i:
                run.add_break()
            run.add_text(line)

        text = self._place(
            self.d.add_paragraph(f"Kôd {number} – {bib.expand(caption)}", style="Caption"))
        text.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return number

    # -- numbered objects -------------------------------------------------

    def eq(self, *items, label: str = None):
        """A displayed, numbered expression. `label` records the number so the text can
        refer to it with `self.ref(label)` and never cite a stale one."""
        self._equation += 1
        if label:
            self.refs[label] = self._equation
        self._place(M.add_equation(self.d, *items, number=self._equation))
        return self._equation

    def ref(self, label: str) -> str:
        """"(3)" - for use inside a sentence: "prema izrazu " + t.ref("saa")."""
        return f"({self.refs[label]})"

    def _next_number(self, counters: dict, label: str = None) -> str:
        """Allocate the next <chapter>.<n> in one counter series, or return the one
        already reserved for `label` by a forward reference. Figures, tables and listings
        each number independently, which is what the template's separate "Sl.", "Tablica"
        and "Kôd" labels mean."""
        if label and label in self.refs:
            return self.refs[label]
        counters[self.chapter] = counters.get(self.chapter, 0) + 1
        number = f"{self.chapter}.{counters[self.chapter]}"
        if label:
            self.refs[label] = number
        return number

    def _next_figure_number(self, label: str = None) -> str:
        return self._next_number(self._figures, label)

    def _next_table_number(self, label: str = None) -> str:
        return self._next_number(self._tables, label)

    def _next_listing_number(self, label: str = None) -> str:
        return self._next_number(self._listings, label)

    def figure(self, path: str, caption: str, label: str = None, width_cm: float = 14.0):
        """A figure with its caption below, numbered <chapter>.<n>, per the template."""
        number = self._next_figure_number(label)

        style = "slika" if "slika" in self._style_names() else None
        holder = self._place(self.d.add_paragraph(style=style))
        holder.alignment = WD_ALIGN_PARAGRAPH.CENTER
        holder.add_run().add_picture(str(path), width=Cm(width_cm))

        text = self._place(
            self.d.add_paragraph(f"Sl. {number} {bib.expand(caption)}", style="Caption"))
        text.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return number

    def figref(self, label: str) -> str:
        """"Sl. 1.2" - usable before the figure itself is placed, which is the normal
        order: a figure is announced in the text and appears just after."""
        return f"Sl. {self._next_figure_number(label)}"

    def table(self, headers, rows, caption: str, label: str = None,
              widths_cm=None):
        """A table with its caption *above* it, which is where the template puts it -
        unlike a figure, whose caption goes below.

        Cells are plain strings; citations in them are expanded like anywhere else. The
        header row is bolded here rather than left to a table style, since the template
        defines none for tables beyond the grid."""
        number = self._next_table_number(label)

        text = self._place(self.d.add_paragraph(
            f"Tablica {number} {bib.expand(caption)}", style="Caption"))
        text.alignment = WD_ALIGN_PARAGRAPH.CENTER

        table = self.d.add_table(rows=1, cols=len(headers))
        if "Table Grid" in self._style_names():
            table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for cell, title in zip(table.rows[0].cells, headers):
            run = cell.paragraphs[0].add_run(bib.expand(str(title)))
            run.bold = True
        for values in rows:
            for cell, value in zip(table.add_row().cells, values):
                cell.paragraphs[0].add_run(bib.expand(str(value)))

        if widths_cm:
            # Column widths in python-docx are per-cell, not per-column: setting the
            # column alone is ignored by Word.
            for row in table.rows:
                for cell, width in zip(row.cells, widths_cm):
                    cell.width = Cm(width)

        self._place(table)
        return number

    def tabref(self, label: str) -> str:
        """"Tablica 4.1" - like `figref`, usable before the table is placed."""
        return f"Tablica {self._next_table_number(label)}"

    def coderef(self, label: str) -> str:
        """"Kôd 5.1" - like `figref`, usable before the listing is placed."""
        return f"Kôd {self._next_listing_number(label)}"


def _suppress_numbering(paragraph):
    """Turn numbering off for one heading that shares a numbered style. A numId of 0 is
    the OOXML way to say "this paragraph opts out"."""
    from docx.oxml.ns import qn
    pPr = paragraph._p.get_or_add_pPr()
    for existing in pPr.findall(qn("w:numPr")):
        pPr.remove(existing)
    off = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0")
    numId = OxmlElement("w:numId"); numId.set(qn("w:val"), "0")
    off.append(ilvl); off.append(numId)
    pPr.append(off)


def base_style(document):
    """Times New Roman 12 pt, 1.5 spacing, as the template requires."""
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    return document


def clear_body(document, from_heading: str, to_heading: str):
    """Remove the existing body between two Heading 1 paragraphs, keeping the front
    matter (title page, sazetak, table of contents) and everything from `to_heading` on.

    Rewriting in place rather than starting a fresh document is deliberate: the template's
    styles, the table-of-contents field and the AI-usage statement all live in the
    original file and would be lost by regenerating it from scratch."""
    paragraphs = document.paragraphs
    start = end = None
    for i, p in enumerate(paragraphs):
        if p.style.name == "Heading 1" and p.text.strip() == from_heading:
            start = i
        elif p.style.name == "Heading 1" and p.text.strip() == to_heading:
            end = i
            break
    if start is None or end is None:
        raise ValueError(f"could not locate '{from_heading}' .. '{to_heading}' in the document")

    body = document.element.body
    for p in paragraphs[start:end]:
        body.remove(p._p)
    return start, end


def capture_heading_numbering(document):
    """Take the `numPr` from the first numbered heading at each level, before the body is
    cleared, so regenerated headings rejoin the same numbering sequence."""
    from docx.oxml.ns import qn
    found = {}
    for p in document.paragraphs:
        name = p.style.name
        if not name.startswith("Heading"):
            continue
        level = int(name.split()[-1])
        if level in found:
            continue
        pPr = p._p.pPr
        numPr = pPr.find(qn("w:numPr")) if pPr is not None else None
        if numPr is not None:
            found[level] = numPr
    return found


def numbering_into_heading_styles(document, numbering, unnumbered=()):
    """Move chapter numbering from individual paragraphs into the Heading styles.

    With the numbering on the style, a heading typed by hand in Word later joins the
    sequence automatically, which is not true when each heading carries its own `numPr`.
    The headings that must stay unnumbered - the abstract pages, Uvod, Zakljucak,
    Literatura - then need an explicit override, since they share the same style: a
    `numId` of 0 switches numbering off for that paragraph only.
    """
    from copy import deepcopy

    from docx.oxml.ns import qn

    for level, numPr in numbering.items():
        style = document.styles[f"Heading {level}"]
        pPr = style.element.get_or_add_pPr()
        for existing in pPr.findall(qn("w:numPr")):
            pPr.remove(existing)
        pPr.append(deepcopy(numPr))

    titles = set(unnumbered)
    for p in document.paragraphs:
        if p.style.name.startswith("Heading") and p.text.strip() in titles:
            _suppress_numbering(p)


def write_bibliography(document, entries, after="Literatura", before=None):
    """Replace the bibliography list with `entries`, keeping the heading in place."""
    heading = marker = None
    for p in document.paragraphs:
        if p.style.name == "Heading 1" and p.text.strip() == after:
            heading = p
        elif heading is not None and p.style.name == "Heading 1":
            marker = p
            break
    if heading is None:
        raise ValueError(f"heading '{after}' not found")

    # Compare the underlying XML elements, not the Paragraph wrappers: python-docx
    # builds fresh wrapper objects on every access, so identity between two reads of
    # document.paragraphs never holds and the old list would survive alongside the new.
    body = document.element.body
    heading_el = heading._p
    marker_el = marker._p if marker is not None else None
    started = False
    for p in list(document.paragraphs):
        if p._p is heading_el:
            started = True
            continue
        if not started:
            continue
        if marker_el is not None and p._p is marker_el:
            break
        body.remove(p._p)

    # The template asks for the "literatura" style on bibliography entries. Its numbering
    # reference is dangling in this document, so it contributes typography without
    # stamping numbers on the entries - which is what author-year citations need, since a
    # number here would refer to nothing in the text.
    style = "literatura" if "literatura" in {s.name for s in document.styles} else None
    anchor = marker._p if marker is not None else None
    for entry in entries:
        paragraph = document.add_paragraph(entry, style=style)
        if anchor is not None:
            anchor.addprevious(paragraph._p)
